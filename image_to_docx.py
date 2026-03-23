"""
image_to_docx.py — Convert text images to Microsoft Word documents using OCR.

Dependencies:
    pip install pytesseract python-docx Pillow

Tesseract OCR must also be installed:
    Windows: https://github.com/UB-Mannheim/tesseract/wiki
    macOS:   brew install tesseract
    Linux:   sudo apt install tesseract-ocr

Usage:
    # Single image
    python image_to_docx.py photo.png

    # Multiple images → one combined .docx
    python image_to_docx.py page1.png page2.jpg page3.tiff

    # Specify output path
    python image_to_docx.py scan.png --output result.docx

    # Specify Tesseract language (default: eng)
    python image_to_docx.py french_doc.png --lang fra
"""

import argparse
import os
import sys

try:
    from PIL import Image
except ImportError:
    sys.exit("Missing dependency: pip install Pillow")

try:
    import pytesseract
except ImportError:
    sys.exit("Missing dependency: pip install pytesseract")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("Missing dependency: pip install python-docx")


# ── Optional: set this if Tesseract is not on your PATH ──────────────────────
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def preprocess_image(image: Image.Image) -> Image.Image:
    """Convert to greyscale to improve OCR accuracy."""
    return image.convert("L")


def extract_text(image_path: str, lang: str) -> str:
    """Run Tesseract OCR on a single image and return the extracted text."""
    img = Image.open(image_path)
    img = preprocess_image(img)
    text = pytesseract.image_to_string(img, lang=lang)
    return text


def build_document(texts: list[tuple[str, str]], output_path: str) -> None:
    """
    Build a .docx from a list of (image_path, extracted_text) pairs.

    Each image gets a heading and its extracted text as body paragraphs.
    Multiple blank lines in the OCR output are collapsed to a single blank line.
    """
    doc = Document()

    # ── Page setup: US Letter, 1-inch margins ────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── Default body style ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Courier New"   # monospace fits OCR-extracted text well
    style.font.size = Pt(11)

    for idx, (image_path, text) in enumerate(texts):
        # Section heading (filename)
        heading = doc.add_heading(os.path.basename(image_path), level=1)
        heading.runs[0].font.name = "Arial"

        if not text.strip():
            doc.add_paragraph("[No text detected]")
        else:
            # Split into lines, collapse consecutive blank lines
            lines = text.splitlines()
            prev_blank = False
            for line in lines:
                stripped = line.strip()
                if stripped == "":
                    if not prev_blank:
                        doc.add_paragraph("")   # single blank separator
                    prev_blank = True
                else:
                    doc.add_paragraph(stripped)
                    prev_blank = False

        # Page break between images (skip after last)
        if idx < len(texts) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"Saved: {output_path}")


def derive_output_path(image_paths: list[str]) -> str:
    """Default output name: stem of first image + .docx"""
    stem = os.path.splitext(os.path.basename(image_paths[0]))[0]
    return f"{stem}.docx"


def main():
    parser = argparse.ArgumentParser(
        description="Convert text image(s) to a Microsoft Word document via OCR."
    )
    parser.add_argument(
        "images",
        nargs="+",
        metavar="IMAGE",
        help="One or more image files (PNG, JPG, TIFF, BMP, etc.)",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        metavar="FILE",
        help="Output .docx path (default: <first_image_stem>.docx)",
    )
    parser.add_argument(
        "--lang", "-l",
        default="eng",
        metavar="LANG",
        help="Tesseract language code, e.g. eng, fra, deu (default: eng)",
    )
    args = parser.parse_args()

    output_path = args.output or derive_output_path(args.images)

    results: list[tuple[str, str]] = []
    for path in args.images:
        if not os.path.isfile(path):
            print(f"Warning: file not found — {path}", file=sys.stderr)
            continue
        print(f"Processing: {path} …", end=" ", flush=True)
        text = extract_text(path, args.lang)
        char_count = len(text.strip())
        print(f"{char_count} characters extracted")
        results.append((path, text))

    if not results:
        sys.exit("No valid images to process.")

    build_document(results, output_path)


if __name__ == "__main__":
    main()
